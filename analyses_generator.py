import io
from datetime import datetime
from typing import List, Dict, Any
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex: str):
    """Устанавливает фоновый цвет ячейки таблицы."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_analyses_docx(
    patient_name: str,
    analyses_data: List[Dict[str, Any]],
    doctor_name: str = "Врач-специалист",
    center_name: str = "Центр ментального здоровья детей «Маленькая Страна»"
) -> bytes:
    """
    Генерирует официальный DOCX-документ хронологии анализов ребенка на базе python-docx.
    Возвращает байтовый поток .docx.
    """
    doc = docx.Document()
    
    # Поля страницы
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # 1. Шапка центра
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_center = p_header.add_run(f"{center_name}\nМедицинская информационная система (цмз.site)")
    run_center.font.size = Pt(9)
    run_center.font.color.rgb = RGBColor(100, 116, 139)

    # 2. Заголовок документа
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    clean_patient = patient_name.replace("disk:/", "").strip()
    run_title = p_title.add_run(f"Хронология анализов пациента: {clean_patient}")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59)

    # 3. Метаданные отчета
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Дата формирования: {now_str}\n").font.size = Pt(10)
    p_meta.add_run(f"Сформировал: {doctor_name}\n").font.size = Pt(10)
    p_meta.add_run(f"Всего показателей в выписке: {len(analyses_data)}").font.size = Pt(10)

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(8)

    # 4. Таблица анализов
    # Колонки: Дата | Анализ / Показатель | Значение | Норма | Отклонение и Динамика | Комментарий
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = [
        "Дата",
        "Анализ / Показатель",
        "Результат",
        "Норма",
        "Отклонение",
        "Комментарий"
    ]
    hdr_widths = [Inches(1.0), Inches(2.0), Inches(1.0), Inches(1.1), Inches(1.1), Inches(1.3)]

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "2563EB") # Синий фон заголовка
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[idx].width = hdr_widths[idx]

    if not analyses_data:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[5])
        row_cells[0].text = "В медицинских документах пациента не обнаружено структурированных лабораторных анализов."
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for item in analyses_data:
            row_cells = table.add_row().cells
            for idx in range(6):
                row_cells[idx].width = hdr_widths[idx]
                row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            is_out = bool(item.get("is_out_of_norm", False))
            is_rep = bool(item.get("is_repeated", False))
            dyn = item.get("dynamics", "")
            
            # Чередование фона или подсветка повторных
            if is_rep:
                for c in row_cells:
                    set_cell_background(c, "F8FAFC")

            # 0. Дата
            row_cells[0].text = str(item.get("date", "-"))
            p_date = row_cells[0].paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_date.runs[0].font.size = Pt(9)

            # 1. Название анализа
            test_title = item.get("test_name", item.get("parameter", "-"))
            row_cells[1].text = test_title
            p_test = row_cells[1].paragraphs[0]
            p_test.runs[0].font.size = Pt(9)
            if is_rep:
                p_test.runs[0].font.bold = True

            # 2. Результат
            row_cells[2].text = str(item.get("value", "-"))
            p_val = row_cells[2].paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_val.runs[0].font.size = Pt(9)
            if is_out:
                p_val.runs[0].font.bold = True
                p_val.runs[0].font.color.rgb = RGBColor(220, 38, 38) # Красный

            # 3. Норма
            row_cells[3].text = str(item.get("norm", "-"))
            p_norm = row_cells[3].paragraphs[0]
            p_norm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_norm.runs[0].font.size = Pt(8.5)

            # 4. Отклонение и Динамика
            dev_text = str(item.get("deviation", "В норме"))
            if dyn:
                dev_text += f" ({dyn})"
            row_cells[4].text = dev_text
            p_dev = row_cells[4].paragraphs[0]
            p_dev.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_dev.runs[0].font.size = Pt(8.5)
            if is_out:
                p_dev.runs[0].font.bold = True
                p_dev.runs[0].font.color.rgb = RGBColor(220, 38, 38)

            # 5. Комментарий
            row_cells[5].text = str(item.get("comment", "-"))
            p_comm = row_cells[5].paragraphs[0]
            p_comm.runs[0].font.size = Pt(8.5)

    # 5. Подвал с подписью
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    p_footer = doc.add_paragraph()
    p_footer.add_run(f"___________________________ / {doctor_name}\n").font.size = Pt(10)
    p_footer.add_run("Документ сформирован автоматически в системе «Маленькая Страна».").font.italic = True
    p_footer.runs[1].font.size = Pt(8.5)
    p_footer.runs[1].font.color.rgb = RGBColor(148, 163, 184)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
